"""Build the combined streamlining demo docx from per-study drafts.

Updated 2026-05-10 second pass:
- Prefers `study_X.draft.merged.md` (Claude + Codex parallel passes merged
  with explicit DIVERGENCES section) when present; falls back to
  `study_X.draft.md` for studies that weren't redrafted in this pass.
- Renders the DIVERGENCES section per study so Jose can see Claude vs. Codex
  alternative wordings inline in the demo.
- Cheat-sheet section now includes the new 2026-05-10 Table 2 reference rule.

Loads each study's draft, then builds a single docx mirroring the layout of
the earlier Streamlining_Demo_Studies_3A_4_5.docx:

  - Title + subtitle
  - Why this companion doc (intro)
  - Cross-study cheat sheet (Katy's red lines + 2026-05-10 amendment)
  - Per-study sections, each with:
      Study X header
      BEFORE block (bolded section labels Methods/Results inline-bold)
      AFTER block (Claude pass primary)
      Word count line (italic)
      DIVERGENCES (Claude vs. Codex alternative wordings) — when present
      Preservation notes (Claude primary; Codex appended for transparency)
  - Final summary table: Section / Before / After / Saved

Run from revision-analysis/. Outputs Streamlining_Demo_Studies_3A_4_5.docx.
"""

from __future__ import annotations
import re
from pathlib import Path
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

HERE = Path(__file__).resolve().parent
BRIEFS_DIR = HERE / "_streamline_briefs"
OUT_DOCX = HERE / "Streamlining_Demo_Studies_3A_4_5.docx"

# Order and display titles for each study section in the deliverable
STUDIES = [
    ("study_3a", "Study 3A: Methods + Results"),
    ("study_3_discussion", "Study 3 Discussion (covers 3A and 3B + supplementary S2A/S2B)"),
    ("study_4a", "Study 4A: Methods + Results + Mediation"),
    ("study_4b", "Study 4B: Methods + Results + Mediation + Pretest"),
    ("study_4_discussion", "Study 4 Discussion (covers 4A and 4B)"),
    ("study_5", "Study 5: Methods + Results"),
    ("study_5_discussion", "Study 5 Discussion"),
]

WORDCOUNT_RE = re.compile(r"Word count:\s*([\d,]+)\s*->\s*([\d,]+)\s*\(saved\s*([\d,]+)\)", re.IGNORECASE)


def parse_draft(text: str) -> dict:
    """Split a draft into BEFORE / AFTER / WC / NOTES (and DIVERGENCES if a
    merged-format file).

    Recognized top-level headers (## ...): BEFORE, AFTER, "Word count: ...",
    "DIVERGENCES (Claude vs. Codex passes)", "Preservation notes (Claude
    pass; primary)", "Preservation notes (Codex pass; secondary, ...)".
    Falls back gracefully when only the legacy 4-section format is present.
    """
    sections = {
        "BEFORE": "",
        "AFTER": "",
        "WC_LINE": "",
        "NOTES": "",
        "DIVERGENCES": "",
        "CODEX_NOTES": "",
    }
    cursor = None
    buf = {"BEFORE": [], "AFTER": [], "NOTES": [], "DIVERGENCES": [], "CODEX_NOTES": []}
    for line in text.splitlines():
        s = line.strip()
        if s.startswith("## BEFORE"):
            cursor = "BEFORE"; continue
        if s.startswith("## AFTER"):
            cursor = "AFTER"; continue
        if s.startswith("## Word count"):
            sections["WC_LINE"] = s.lstrip("# ").strip()
            cursor = None
            continue
        if s.startswith("## DIVERGENCES"):
            cursor = "DIVERGENCES"; continue
        if s.startswith("## Preservation notes (Claude") or (
            s.startswith("## Preservation notes") and not sections["NOTES"] and cursor != "CODEX_NOTES"
        ):
            cursor = "NOTES"; continue
        if s.startswith("## Preservation notes (Codex"):
            cursor = "CODEX_NOTES"; continue
        if s.startswith("## "):
            cursor = None; continue
        if cursor:
            buf[cursor].append(line)
    for k in ("BEFORE", "AFTER", "NOTES", "DIVERGENCES", "CODEX_NOTES"):
        sections[k] = "\n".join(buf[k]).strip("\n")
    # Parse the word counts numerically too
    m = WORDCOUNT_RE.search(sections["WC_LINE"])
    if m:
        sections["WC_BEFORE"] = int(m.group(1).replace(",", ""))
        sections["WC_AFTER"] = int(m.group(2).replace(",", ""))
        sections["WC_SAVED"] = int(m.group(3).replace(",", ""))
    else:
        sections["WC_BEFORE"] = sections["WC_AFTER"] = sections["WC_SAVED"] = 0
    return sections


def find_draft(study_id: str) -> Path | None:
    """Prefer the merged file when available, then claude, then legacy."""
    for suffix in (".draft.merged.md", ".draft.claude.md", ".draft.md"):
        p = BRIEFS_DIR / f"{study_id}{suffix}"
        if p.exists():
            return p
    return None


def add_para(doc, text="", bold=False, italic=False, size=11, style=None, align=None):
    p = doc.add_paragraph(style=style) if style else doc.add_paragraph()
    if align is not None:
        p.alignment = align
    if text:
        run = p.add_run(text)
        run.bold = bold
        run.italic = italic
        run.font.size = Pt(size)
    return p


def add_runs(doc, runs, style=None, align=None):
    """Add a paragraph composed of multiple runs.
    runs: list of (text, {bold, italic, size, color}) tuples.
    """
    p = doc.add_paragraph(style=style) if style else doc.add_paragraph()
    if align is not None:
        p.alignment = align
    for txt, opts in runs:
        r = p.add_run(txt)
        r.bold = bool(opts.get("bold"))
        r.italic = bool(opts.get("italic"))
        size = opts.get("size", 11)
        r.font.size = Pt(size)
        col = opts.get("color")
        if col:
            r.font.color.rgb = RGBColor(*col)
    return p


# Renderers for BEFORE/AFTER blocks. The Codex source is markdown-ish text:
# study heading, then 'Methods', 'Participants.', 'Procedure.', 'Results',
# etc., with the bolded labels inline. We render each non-empty line as a
# paragraph; lines that are exactly "Methods"/"Results"/"Discussion"/study
# heading get bold treatment; lines starting with "Participants." or
# "Procedure." get bolded leading label + remainder as plain text.
SECTION_HEADERS = {"Methods", "Results", "Discussion", "Mediation", "Pretest"}
LABEL_PREFIXES = (
    "Participants.",
    "Procedure.",
    "Mediation.",
    "Pretest.",
    "Internal motivation to respond without prejudice.",
    "External motivation to respond without prejudice.",
    "Political ideology.",
    "Political party affiliation.",
    "Mechanism measures.",
    "Mediator and moderator measures.",
    "Exploratory Mediation Analyses.",
    "Fairness concerns.",
)


def render_block(doc, body: str):
    for raw in body.splitlines():
        line = raw.rstrip()
        if not line.strip():
            continue
        stripped = line.strip()
        # Section headers
        if stripped in SECTION_HEADERS or re.match(r"^Study\s+[0-9]+[A-Z]?\b", stripped):
            add_para(doc, stripped, bold=True, size=12)
            continue
        # "Label. rest..." pattern
        for prefix in LABEL_PREFIXES:
            if stripped.startswith(prefix):
                rest = stripped[len(prefix):]
                add_runs(doc, [
                    (prefix, {"bold": True, "size": 11}),
                    (rest, {"size": 11}),
                ])
                break
        else:
            # Default plain paragraph
            add_para(doc, stripped, size=11)


def render_notes(doc, notes_md: str):
    for raw in notes_md.splitlines():
        line = raw.rstrip()
        if not line.strip():
            continue
        stripped = line.strip()
        # Treat lines starting with "- " or "* " as bullets
        if stripped.startswith(("- ", "* ")):
            content = stripped[2:].strip()
            try:
                add_para(doc, content, size=10, style="List Bullet")
            except KeyError:
                # Style may not exist in fresh docx; fall back to plain bullet glyph
                add_para(doc, "• " + content, size=10)
        else:
            add_para(doc, stripped, italic=True, size=10)


def render_divergences(doc, div_md: str):
    """Render the DIVERGENCES block.

    The merger writes paragraphs separated by blank lines, with `### Divergence
    N (...)` headers and `**Claude:** ...` / `**Codex:** ...` labels. We render
    each divergence as: bold heading, then a "Claude:"-prefixed paragraph,
    then a "Codex:"-prefixed paragraph.
    """
    if not div_md:
        return
    current_heading = None
    current_label = None
    para_buf: list[str] = []

    def flush():
        nonlocal para_buf, current_label
        if not para_buf:
            return
        joined = " ".join(s.strip() for s in para_buf if s.strip())
        if not joined:
            para_buf = []
            current_label = None
            return
        if current_label == "Claude":
            add_runs(doc, [
                ("Claude: ", {"bold": True, "size": 10}),
                (joined, {"size": 10}),
            ])
        elif current_label == "Codex":
            add_runs(doc, [
                ("Codex: ", {"bold": True, "size": 10}),
                (joined, {"size": 10}),
            ])
        else:
            add_para(doc, joined, size=10, italic=True)
        para_buf = []

    for raw in div_md.splitlines():
        line = raw.rstrip()
        s = line.strip()
        if not s:
            flush()
            continue
        if s.startswith("### "):
            flush()
            add_para(doc, s[4:].strip(), bold=True, size=11)
            current_label = None
            continue
        if s.startswith("**Claude:**"):
            flush()
            current_label = "Claude"
            content = s[len("**Claude:**"):].strip()
            if content:
                para_buf.append(content)
            continue
        if s.startswith("**Codex:**"):
            flush()
            current_label = "Codex"
            content = s[len("**Codex:**"):].strip()
            if content:
                para_buf.append(content)
            continue
        para_buf.append(s)
    flush()


def add_summary_table(doc, rows):
    """rows: list of (label, before, after, saved)."""
    table = doc.add_table(rows=1 + len(rows), cols=4)
    table.style = "Light List Accent 1"
    hdr = table.rows[0].cells
    for i, name in enumerate(("Section", "Before (words)", "After (words)", "Saved")):
        hdr[i].paragraphs[0].add_run(name).bold = True
    total_b = total_a = total_s = 0
    for i, (label, b, a, s) in enumerate(rows, start=1):
        cells = table.rows[i].cells
        cells[0].text = label
        cells[1].text = str(b)
        cells[2].text = str(a)
        cells[3].text = str(s)
        total_b += b; total_a += a; total_s += s
    # Totals row
    row = table.add_row().cells
    row[0].paragraphs[0].add_run("TOTAL").bold = True
    row[1].paragraphs[0].add_run(str(total_b)).bold = True
    row[2].paragraphs[0].add_run(str(total_a)).bold = True
    row[3].paragraphs[0].add_run(str(total_s)).bold = True
    return total_b, total_a, total_s


def main():
    doc = Document()
    # Page margin shrink for more density
    for sec in doc.sections:
        sec.left_margin = Inches(1.0)
        sec.right_margin = Inches(1.0)
        sec.top_margin = Inches(0.9)
        sec.bottom_margin = Inches(0.9)

    # Title
    add_para(doc, "Streamlining Demo: Studies 3A, 4, and 5 — second pass", bold=True, size=16)
    add_para(doc, "Manuscript: Does Feedback Enhance Diversity in Selection Decisions?  |  Editor: George Wu (MS R2)  |  Date: 2026-05-10", italic=True, size=10)
    add_para(doc)

    # Why this doc
    add_para(doc, "Why this companion doc", bold=True, size=13)
    add_para(doc,
             "This is the SECOND streamlining pass on Studies 3A, 4 (4A and 4B), and 5 of the manuscript. The first pass produced AFTER drafts that Katy reverted in places as too aggressive; this second pass (a) extracts implicit rules from the accepted post-Katy text in Studies 1, 2, and 3B (codified in RULES_DISCOVERED.md), (b) leverages the new cross-study Table 2 to compress the verbose Wald F-test enumeration in Studies 3A, 4A, and 4B to a single sentence + Table 2 pointer (Study 1 keeps the verbose form as place of first mention; Study 5 keeps its 2x2 Wald z-test as study-specific), and (c) ran two independent drafting passes (Claude with consistency emphasis, Codex with holistic-flow emphasis) and surfaces every divergence between them so Jose can resolve in his manual review.",
             italic=True, size=10)
    add_para(doc)

    # Cross-study Katy red lines
    add_para(doc, "Cross-study constraints (Katy's red lines plus 2026-05-10 amendment)", bold=True, size=13)
    redlines = [
        "Replication-grade detail must survive somewhere. Demographics, inclusion thresholds, balance-check stats, procedural display details. If repeated, a single footnote on Study 1 covers later studies (\"unless otherwise stated\").",
        "Detailed write-up appears once in Study 1; later studies condense. Study 1 stays verbose.",
        "DVs must be named at every report (\"chose a biopic featuring a woman\"), not collapsed to bare \"%\".",
        "All key stats appear somewhere. \"all p's > X\" is allowed in main text only when an appendix carries the full numbers.",
        "Original wording often beats rewrite. Don't replace clearer original prose just to compress.",
        "2026-05-10 amendment: With the new cross-study Table 2 in place, Studies 3A, 4A, and 4B replace the verbose `(1)... (2)... (3)...` Wald F-test enumeration with a single sentence + Table 2 pointer. Study 1 keeps the verbose enumeration. Study 5 is not in Table 2 and keeps its own analytical structure (2x2 reversal Wald z-test).",
        "Mediation in Studies 4A and 4B retained in full detail (Sobel, ACME, multiple mediation, alpha values). It's the contribution.",
        "Pretest in Study 4B condensed to 1-2 sentences with appendix pointer.",
        "Study 5 RETAINS pooled interaction analysis, 2x2 design, fairness-concerns mediator, separate underrep / overrep regressions.",
        "Newly discovered rule: balance-check sentence in Studies with clean balance compresses to baseline percentage + \"with balance across conditions (p = X; see Appendix Table SX)\" form (matching post-Katy Study 3B). Study 5's 2x2 design retains per-cell t/p/CI.",
        "Newly discovered rule: avoid em/en dashes in newly authored prose (per standing user preference). Use commas, parentheses, or sentence breaks instead.",
    ]
    for r in redlines:
        try:
            add_para(doc, r, size=10, style="List Bullet")
        except KeyError:
            add_para(doc, "• " + r, size=10)
    add_para(doc)

    # Per-study sections
    summary_rows = []
    for study_id, display_title in STUDIES:
        draft_path = find_draft(study_id)
        if draft_path is None:
            print(f"[WARN] no draft found for {study_id}")
            continue
        sections = parse_draft(draft_path.read_text(encoding="utf-8"))

        # Section break header
        doc.add_page_break()
        add_para(doc, display_title, bold=True, size=15)
        add_para(doc, f"Source draft: {draft_path.name}", italic=True, size=9)
        add_para(doc)

        # BEFORE
        b_label = f"BEFORE  (current draft, {sections['WC_BEFORE']} words)"
        add_para(doc, b_label, bold=True, size=12)
        render_block(doc, sections["BEFORE"])
        add_para(doc)

        # AFTER
        saved = sections["WC_SAVED"]
        a_label_words = "saves" if saved >= 0 else "adds"
        a_label = f"AFTER  (streamlined, {sections['WC_AFTER']} words; {a_label_words} {abs(saved)} words)"
        add_para(doc, a_label, bold=True, size=12)
        render_block(doc, sections["AFTER"])
        add_para(doc)

        # Word-count line (already inside AFTER label, but keep explicit)
        add_para(doc, sections["WC_LINE"], italic=True, size=10)
        add_para(doc)

        # DIVERGENCES (only present in merged drafts)
        if sections.get("DIVERGENCES"):
            add_para(doc, "Divergences (Claude vs. Codex)", bold=True, size=12)
            render_divergences(doc, sections["DIVERGENCES"])
            add_para(doc)

        # Preservation notes (Claude pass primary)
        notes_label = "Preservation notes (Claude pass; primary)" if sections.get("CODEX_NOTES") else "Preservation notes"
        add_para(doc, notes_label, bold=True, size=12)
        render_notes(doc, sections["NOTES"])
        add_para(doc)

        # Codex preservation notes (transparency, when present)
        if sections.get("CODEX_NOTES"):
            add_para(doc, "Preservation notes (Codex pass; secondary, for transparency)", bold=True, size=12)
            render_notes(doc, sections["CODEX_NOTES"])
            add_para(doc)

        summary_rows.append((display_title, sections["WC_BEFORE"], sections["WC_AFTER"], sections["WC_SAVED"]))

    # Final summary table
    doc.add_page_break()
    add_para(doc, "Word-count savings tally", bold=True, size=15)
    if summary_rows:
        total_b, total_a, total_s = add_summary_table(doc, summary_rows)
        add_para(doc)
        pages = total_s / 250.0  # rough double-spaced page estimate
        add_para(doc,
                 f"Total saved: {total_s} words across {len(summary_rows)} sections "
                 f"(~{pages:.1f} double-spaced pages at 250 wpm).",
                 italic=True, size=10)

    OUT_DOCX.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUT_DOCX)
    print(f"wrote {OUT_DOCX} ({OUT_DOCX.stat().st_size} bytes)")


if __name__ == "__main__":
    main()
