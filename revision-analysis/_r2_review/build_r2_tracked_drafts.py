"""Build Manuscript_R2_drafts_TRACKED.docx by injecting Word tracked-change revisions
into the freshly exported manuscript Gdoc.

Inputs (must already exist; produced by fetch_manuscript_gdoc.py + fetch_response_letter_gdoc.py):
  - manuscript_gdoc_export.docx  (Drive export of the live manuscript)
  - manuscript_gdoc_text.txt     (plain-text walk for anchor validation)
  - response_letter_gdoc_text.txt (canonical promise wording)

Output:
  - Manuscript_R2_drafts_TRACKED.docx (copy of the export with w:ins / w:del revisions)
  - build_r2_tracked_drafts.log (run log written by tee from the calling shell)

The script validates:
  1. All required inputs exist and are non-empty.
  2. Every anchor opening phrase appears in manuscript_gdoc_text.txt (drift check).
  3. After save, the output .docx reopens with python-docx and contains the expected
     count of w:ins elements with the right author string.

Drafts implemented (specs match the plan file):
  1. Intro feedback/goals citation refresh (inline)
  2. Intro implicit-injunctive-norm framing (new paragraph)
  3. GD foundational-citation refresh (inline)
  4. GD two-pathway framing (new paragraph)
  5A. Study 4 Discussion political-ideology null sentence (append at paragraph end)
  5B. GD two-pathway paragraph closing sentence (added inline to Draft 4)
  6. Study 1 Discussion NPR-importance pretest (new paragraph)
  7. References bibliography additions (new paragraphs at alphabetical positions)

Out of scope (not drafted here):
  - Study 4B / Table R1 zero-initial footnote
  - R2 first-comment socio-political climate paragraphs
"""
from __future__ import annotations

import shutil
import sys
from pathlib import Path

from docx import Document
from lxml import etree


# --- Paths ------------------------------------------------------------------

ROOT = Path(r"C:/Users/jcerv/Jose/diversity_feedback/revision-analysis/_r2_review")
SRC_DOCX = ROOT / "manuscript_gdoc_export.docx"
SRC_TEXT = ROOT / "manuscript_gdoc_text.txt"
LETTER_TEXT = ROOT / "response_letter_gdoc_text.txt"
OUT_DOCX = ROOT / "Manuscript_R2_drafts_TRACKED.docx"

AUTHOR = "Claude (R2 draft)"
DATE = "2026-05-12T00:00:00Z"

W_NS = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
W = "{%s}" % W_NS
NSMAP = {"w": W_NS}


# --- Anchor specs -----------------------------------------------------------
#
# Each anchor is a prefix that uniquely identifies a manuscript paragraph.
# Validation step confirms each prefix appears in manuscript_gdoc_text.txt.

ANCHORS = {
    "draft1_p9_intro_FIT": (
        "According to Feedback Intervention Theory, feedback most effectively regulates"
    ),
    "draft1_p10_intro_Schultz": (
        "Just as feedback accompanied by goals can motivate behavior change"
    ),
    "draft2_p17_intro_We_propose": (
        "We propose that when descriptive feedback reveals that an evaluator"
    ),
    "draft3_p170_GD_FIT": (
        "Our work makes two primary theoretical contributions. First, we offer new insights"
    ),
    "draft5A_p132_study4_pol_ideology": (
        "In addition, Study 4B established that the effect of gender feedback was not moderated"
    ),
    "draft6_p50_study1_discussion": (
        "Study 1 mirrored a real organizational practice currently deployed by NPR"
    ),
}


# --- Helpers ----------------------------------------------------------------


def _q(name: str) -> str:
    return W + name


def make_run(text: str, *, bold: bool = False, italic: bool = False) -> etree._Element:
    r = etree.SubElement(etree.Element("dummy"), _q("r"))
    if bold or italic:
        rPr = etree.SubElement(r, _q("rPr"))
        if bold:
            etree.SubElement(rPr, _q("b"))
        if italic:
            etree.SubElement(rPr, _q("i"))
    t = etree.SubElement(r, _q("t"))
    t.text = text
    t.set("{http://www.w3.org/XML/1998/namespace}space", "preserve")
    # Detach from dummy parent so it can be appended elsewhere.
    r.getparent().remove(r)
    return r


def make_ins(*, ins_id: int) -> etree._Element:
    ins = etree.Element(_q("ins"), nsmap=NSMAP)
    ins.set(_q("id"), str(ins_id))
    ins.set(_q("author"), AUTHOR)
    ins.set(_q("date"), DATE)
    return ins


def make_del(*, del_id: int) -> etree._Element:
    d = etree.Element(_q("del"), nsmap=NSMAP)
    d.set(_q("id"), str(del_id))
    d.set(_q("author"), AUTHOR)
    d.set(_q("date"), DATE)
    return d


def make_delrun(text: str) -> etree._Element:
    """A <w:r> whose visible text is a <w:delText> (deletion marker)."""
    r = etree.Element(_q("r"))
    dt = etree.SubElement(r, _q("delText"))
    dt.text = text
    dt.set("{http://www.w3.org/XML/1998/namespace}space", "preserve")
    return r


_ins_counter = [1000]


def next_ins_id() -> int:
    _ins_counter[0] += 1
    return _ins_counter[0]


def new_paragraph_with_ins(text: str, *, pPr_source: etree._Element | None = None) -> etree._Element:
    """Build a <w:p> whose entire content is wrapped in <w:ins>."""
    p = etree.Element(_q("p"), nsmap=NSMAP)
    if pPr_source is not None:
        pPr_clone = etree.fromstring(etree.tostring(pPr_source))
        p.append(pPr_clone)
    else:
        pPr = etree.SubElement(p, _q("pPr"))
    ins = make_ins(ins_id=next_ins_id())
    ins.append(make_run(text))
    p.append(ins)
    return p


def find_paragraph_by_prefix(doc: Document, prefix: str):
    pref = prefix.strip()
    for p in doc.paragraphs:
        txt = p.text or ""
        if txt.strip().startswith(pref):
            return p
    return None


def append_inline_ins(paragraph, text: str) -> None:
    """Append an inserted run at the end of an existing paragraph."""
    ins = make_ins(ins_id=next_ins_id())
    ins.append(make_run(text))
    paragraph._p.append(ins)


def replace_text_in_paragraph(paragraph, find_str: str, replace_str: str) -> bool:
    """Find find_str in the paragraph's combined run text; split runs, wrap the matched
    span in <w:del> and append <w:ins> with replace_str after it.

    Returns True if a substitution was applied; False if find_str was not found.

    Limitation: this implementation operates on the first contiguous match across
    runs and assumes find_str lives entirely within the visible text (not inside
    field codes, hyperlinks, etc.). Works for ordinary body paragraphs.
    """
    runs = paragraph.runs
    if not runs:
        return False

    # Build per-run text positions.
    positions = []
    full = ""
    for run in runs:
        start = len(full)
        full += run.text or ""
        positions.append((run, start, len(full)))

    idx = full.find(find_str)
    if idx == -1:
        return False
    end = idx + len(find_str)

    # Find which runs (and offsets within them) cover [idx, end).
    affected = []  # list of (run, local_start, local_end)
    for run, rs, re_ in positions:
        if re_ <= idx or rs >= end:
            continue
        local_start = max(0, idx - rs)
        local_end = min(re_ - rs, end - rs)
        affected.append((run, local_start, local_end))

    if not affected:
        return False

    # Split the first and last affected runs so the deleted text occupies its own
    # runs we can wrap in <w:del>.
    new_run_elements = []  # ordered list of (etree _Element, "prefix"|"deleted"|"suffix")

    for i, (run, ls, le) in enumerate(affected):
        text = run.text or ""
        prefix_text = text[:ls]
        del_text = text[ls:le]
        suffix_text = text[le:]
        # Clone original run xml to preserve formatting.
        original_xml = run._element

        if prefix_text:
            r_prefix = etree.fromstring(etree.tostring(original_xml))
            _set_run_text(r_prefix, prefix_text)
            new_run_elements.append((r_prefix, "prefix"))
        if del_text:
            # Build a <w:r> with <w:delText> using original formatting.
            r_del = etree.fromstring(etree.tostring(original_xml))
            _replace_t_with_delText(r_del, del_text)
            new_run_elements.append((r_del, "deleted"))
        if suffix_text and i == len(affected) - 1:
            r_suffix = etree.fromstring(etree.tostring(original_xml))
            _set_run_text(r_suffix, suffix_text)
            new_run_elements.append((r_suffix, "suffix"))

        # Remove the original run from the paragraph; we will reinsert in order.
        original_xml.getparent().remove(original_xml)

    # Re-insert the elements where the first affected run lived. We track that by
    # inserting them at the end of the paragraph in order — but for the formatting
    # of the paragraph that is acceptable here because all affected runs are
    # contiguous text-only runs with no neighbours we removed in between.
    # However, for correctness we instead insert at the original position of the
    # first affected run. The original anchor was just removed, so we rely on the
    # ordering of remaining runs. To avoid order corruption, we collected affected
    # runs above in document order, and now we append in document order; for a
    # single-paragraph inline tweak (the only case we use this for) this matches
    # the original position.

    # Wrap deletion runs in a single <w:del>; place <w:ins> with the new content
    # immediately after.
    del_elem = make_del(del_id=next_ins_id())
    ins_elem = make_ins(ins_id=next_ins_id())
    ins_elem.append(make_run(replace_str))

    # Build the final ordered list of new XML elements to append at the
    # paragraph end (assumes only one contiguous match and only-text runs were
    # affected, which is true for our use cases).
    pPr_first_child = paragraph._p.find(_q("pPr"))
    insertion_anchor = pPr_first_child if pPr_first_child is not None else None

    seq = []
    deletion_open = False
    for r_elem, kind in new_run_elements:
        if kind == "deleted":
            if not deletion_open:
                seq.append(("open_del", None))
                deletion_open = True
            seq.append(("inside_del", r_elem))
        else:
            if deletion_open:
                seq.append(("close_del", None))
                seq.append(("ins_with_replacement", None))
                deletion_open = False
            seq.append(("normal", r_elem))
    if deletion_open:
        seq.append(("close_del", None))
        seq.append(("ins_with_replacement", None))

    cursor = insertion_anchor  # we insert after this; if None, append to paragraph.
    current_del = None
    for kind, payload in seq:
        if kind == "normal":
            _append_after(paragraph._p, cursor, payload)
            cursor = payload
        elif kind == "open_del":
            current_del = make_del(del_id=next_ins_id())
            _append_after(paragraph._p, cursor, current_del)
            cursor = current_del
        elif kind == "inside_del":
            current_del.append(payload)
        elif kind == "close_del":
            current_del = None
        elif kind == "ins_with_replacement":
            ins_node = make_ins(ins_id=next_ins_id())
            ins_node.append(make_run(replace_str))
            _append_after(paragraph._p, cursor, ins_node)
            cursor = ins_node

    return True


def _append_after(parent: etree._Element, anchor: etree._Element | None, child: etree._Element) -> None:
    if anchor is None:
        parent.append(child)
    else:
        anchor.addnext(child)


def _set_run_text(r_elem: etree._Element, text: str) -> None:
    # Remove existing <w:t> children and add a fresh one.
    for t in r_elem.findall(_q("t")):
        r_elem.remove(t)
    t = etree.SubElement(r_elem, _q("t"))
    t.text = text
    t.set("{http://www.w3.org/XML/1998/namespace}space", "preserve")


def _replace_t_with_delText(r_elem: etree._Element, text: str) -> None:
    for t in r_elem.findall(_q("t")):
        r_elem.remove(t)
    for dt in r_elem.findall(_q("delText")):
        r_elem.remove(dt)
    dt = etree.SubElement(r_elem, _q("delText"))
    dt.text = text
    dt.set("{http://www.w3.org/XML/1998/namespace}space", "preserve")


# --- Draft prose ------------------------------------------------------------

DRAFT1_INTRO_CITATION_FIND = "(Kluger & DeNisi, 1996, 1998)"
DRAFT1_INTRO_CITATION_REPLACE = (
    "(Kluger & DeNisi, 1996, 1998; Locke & Latham, 2019; Itzchakov & Latham, 2018; "
    "Chen, Latham, & Piccolo, 2020)"
)

DRAFT1_INTRO_SCHULTZ_APPEND = (
    " A growing applied literature has since shown that descriptive-feedback-plus-norm "
    "designs of this kind change household behavior across countries and settings "
    "(Schultz, Nolan, & Cialdini, 2018; Andor, Gerster, & Peters, 2020; Henry, Ferraro, "
    "& Kontoleon, 2019; Mertens & Schultz, 2021)."
)

DRAFT2_NEW_PARAGRAPH = (
    "Consistent with the broader logic of the feedback literature, feedback regulates "
    "behavior only when there is a standard against which past performance can be "
    "evaluated. The classical literature establishes that pairing feedback with an "
    "explicit goal supplies that standard. But the same logic implies that descriptive "
    "feedback alone can be effective in domains where a strong implicit standard is "
    "already entrenched, such that no explicit goal need be paired alongside it. The "
    "treatment of women and racial minorities in selection decisions provides a clean "
    "test of this argument: anti-discrimination law, social and reputational pressures "
    "to behave in ways that are not prejudicial, and selectors’ personally held "
    "egalitarian self-image jointly sustain an entrenched injunctive norm against "
    "prejudiced selection (Plant & Devine, 1998; Crandall & Eshleman, 2003; "
    "Álvarez-Benjumea, 2023). When descriptive feedback reveals a discrepancy "
    "between this implicit standard and past behavior, we theorize that motivation to "
    "respond without prejudice is activated even in the absence of a specific external "
    "goal."
)

DRAFT3_GD_CITATION_FIND = "(Kluger & DeNisi, 1996; Schultz et al., 2007)"
DRAFT3_GD_CITATION_REPLACE = (
    "(Kluger & DeNisi, 1996; Locke & Latham, 2019; Schultz et al., 2007; "
    "Schultz, Nolan, & Cialdini, 2018)"
)

DRAFT4_NEW_PARAGRAPH = (
    "It is worth explicitly distinguishing two pathways through which descriptive "
    "feedback can operate in this account. The first pathway, the external motivation "
    "to control prejudice, is driven both by a firm’s stated norms and, more "
    "broadly, by the fact that selection decisions are typically observable to "
    "colleagues, prospective candidates, supervisors, and outside stakeholders; the "
    "implicit norm against appearing prejudiced is sustained by anti-discrimination "
    "law, employee and customer scrutiny, and broader reputational concerns. "
    "Organizations’ explicit DEI commitments are themselves often responses to "
    "these external pressures rather than the source of internal organizational norms "
    "(Levi & Fried, 2024); a given firm’s retreat from DEI rhetoric therefore "
    "reflects shifting external pressure on the firm rather than a collapse of the "
    "broader social pressure on individual selection decisions. The second pathway, "
    "the internal motivation to control prejudice, refers to selectors’ "
    "personally held egalitarian beliefs and self-image as non-prejudiced decision "
    "makers. This pathway should not vary with organizational climate, and it is the "
    "pathway most directly supported by Study 4, in which feedback was delivered "
    "privately and internal motivation to respond without prejudice emerged as the "
    "dominant mediator. Empirically, the absence of a political-ideology interaction "
    "in Study 4 is consistent with both pathways operating across partisan lines, and "
    "converges with a recent field experiment in which the effect of demographic-"
    "identity disclosure on city-counselor responsiveness did not differ between "
    "Republican and Democratic counselors (Kirgios et al., 2022). We expect our effect "
    "to generalize wherever selections are observable and selectors hold at least some "
    "egalitarian commitments, and to be somewhat attenuated in fully private or "
    "anonymous selection decisions or in contexts that explicitly encourage "
    "discriminatory behavior."
)

DRAFT5A_APPEND = (
    " Although we are cautious about over-interpreting a null interaction, the pattern "
    "is suggestive that descriptive feedback operates across partisan lines, consistent "
    "with the broader claim that the effect is sustained by injunctive norms shared "
    "across the political spectrum rather than by any single political community’s "
    "stated DEI commitments. This interpretation also converges with a recent field "
    "experiment in which women and racial minorities benefited from explicitly stating "
    "their identity when seeking help from city counselors regardless of whether they "
    "contacted Republicans or Democrats, mediated by counselors’ motivation to "
    "control prejudice (Kirgios et al., 2022)."
)

DRAFT6_NEW_PARAGRAPH = (
    "To address a concern that participants might have been less swayed by descriptive "
    "feedback about the three comparison attributes simply because they viewed those "
    "attributes as less important to NPR’s mission than gender, we ran a between-"
    "subjects pretest with N = 300 Prolific participants. Each participant was "
    "randomly assigned to rate, on a 1 to 7 importance scale, one of the four expert-"
    "attribute statements used in Study 1’s feedback manipulation, with the NPR "
    "producer framing held constant. If anything, the comparison attributes were rated "
    "as at least as important to NPR’s mission as gender: whether a guest was a "
    "woman (M = 3.64, SD = 1.76, n = 76) was rated as significantly less important "
    "than whether a guest worked at a university (M = 4.67, SD = 1.60; t(148.8) = "
    "−3.76, p < .001, d = −0.61) and was under 50 years old (M = 4.44, SD = "
    "1.73; t(149.0) = −2.81, p = .006, d = −0.46), and as statistically "
    "indistinguishable from whether a guest was based on the West Coast of the United "
    "States (M = 3.37, SD = 1.75; t(149.0) = 0.95, p = .343, d = +0.16). Differences "
    "in perceived importance therefore cannot account for why only feedback about "
    "gender moved subsequent selections in Study 1, consistent with our claim that "
    "descriptive feedback motivates behavior specifically on dimensions carrying an "
    "implicit injunctive norm."
)

# Draft 7: References to add (alphabetical APA entries).
DRAFT7_NEW_REFERENCES: list[str] = [
    (
        "Andor, M. A., Gerster, A., & Peters, J. (2020). Social norms and energy "
        "conservation beyond the US. Journal of Environmental Economics and "
        "Management, 103, 102351. https://doi.org/10.1016/j.jeem.2020.102351"
    ),
    (
        "Chen, X., Latham, G. P., & Piccolo, R. F. (2020). An enumerative review and "
        "a meta-analysis of primed goal effects on organizational behavior. Applied "
        "Psychology, 70(1), 216-253. https://doi.org/10.1111/apps.12239"
    ),
    (
        "Devine, P. G., & Ash, T. L. (2022). Diversity training goals, limitations, "
        "and promise: A review of the multidisciplinary literature. Annual Review of "
        "Psychology, 73(1), 403-429. https://doi.org/10.1146/annurev-psych-060221-122215"
    ),
    (
        "Henry, M. L., Ferraro, P. J., & Kontoleon, A. (2019). The behavioural effect "
        "of electronic home energy reports: Evidence from a randomised field trial in "
        "the United States. Energy Policy, 132, 1256-1261. "
        "https://doi.org/10.1016/j.enpol.2019.06.039"
    ),
    (
        "Itzchakov, G., & Latham, G. P. (2018). The moderating effect of performance "
        "feedback and the mediating effect of self-set goals on the primed goal-"
        "performance relationship. Applied Psychology, 69(2), 379-414. "
        "https://doi.org/10.1111/apps.12176"
    ),
    (
        "Levi, A., & Fried, Y. (2024). Diversity, equity, and inclusion programs’ "
        "emphasis on symbolism: Causes and consequences. Journal of Organizational "
        "Behavior, 46(1), 172-187. https://doi.org/10.1002/job.2834"
    ),
    (
        "Locke, E. A., & Latham, G. P. (2019). The development of goal setting "
        "theory: A half century retrospective. Motivation Science, 5(2), 93-105. "
        "https://doi.org/10.1037/mot0000127"
    ),
    (
        "Mertens, S. N., & Schultz, P. W. (2021). Referent group specificity: "
        "Optimizing normative feedback to increase residential recycling. Journal of "
        "Environmental Psychology, 73, 101541. "
        "https://doi.org/10.1016/j.jenvp.2020.101541"
    ),
    (
        "Schultz, P. W., Nolan, J. M., & Cialdini, R. B. (2018). The constructive, "
        "destructive, and reconstructive power of social norms: Reprise. Perspectives "
        "on Psychological Science, 13(2), 249-254. "
        "https://doi.org/10.1177/1745691617693325"
    ),
]


# --- Drift validation -------------------------------------------------------


def validate_inputs() -> None:
    missing = [p for p in (SRC_DOCX, SRC_TEXT, LETTER_TEXT) if not p.exists() or p.stat().st_size == 0]
    if missing:
        raise SystemExit(f"FAIL: missing/empty inputs: {missing}")


def validate_anchors() -> None:
    text = SRC_TEXT.read_text(encoding="utf-8")
    not_found = [k for k, prefix in ANCHORS.items() if prefix not in text]
    if not_found:
        raise SystemExit(
            "FAIL: anchor drift detected; the following keys do not match manuscript text:\n  "
            + "\n  ".join(not_found)
        )


def validate_no_dashes_in_drafts() -> None:
    bad = []
    drafts = {
        "DRAFT1_REPLACE": DRAFT1_INTRO_CITATION_REPLACE,
        "DRAFT1_APPEND": DRAFT1_INTRO_SCHULTZ_APPEND,
        "DRAFT2": DRAFT2_NEW_PARAGRAPH,
        "DRAFT3_REPLACE": DRAFT3_GD_CITATION_REPLACE,
        "DRAFT4": DRAFT4_NEW_PARAGRAPH,
        "DRAFT5A": DRAFT5A_APPEND,
        "DRAFT6": DRAFT6_NEW_PARAGRAPH,
    }
    for name, s in drafts.items():
        if "–" in s or "—" in s:
            bad.append(name)
        if "latent" in s.lower():
            bad.append(f"{name}:contains 'latent'")
    if bad:
        raise SystemExit("FAIL: voice-constraint violations in drafts: " + ", ".join(bad))


# --- Apply drafts -----------------------------------------------------------


def apply_drafts(doc: Document) -> None:
    # ---- Draft 1: Intro citation refresh + Schultz applied-literature append ----
    p_fit = find_paragraph_by_prefix(doc, ANCHORS["draft1_p9_intro_FIT"])
    if p_fit is None:
        raise SystemExit("FAIL: Draft 1 anchor 1 (FIT paragraph) not found")
    if not replace_text_in_paragraph(p_fit, DRAFT1_INTRO_CITATION_FIND, DRAFT1_INTRO_CITATION_REPLACE):
        raise SystemExit("FAIL: Draft 1 anchor 1 inline replace failed")
    print("OK Draft 1 (Intro FIT citations refreshed)")

    p_schultz = find_paragraph_by_prefix(doc, ANCHORS["draft1_p10_intro_Schultz"])
    if p_schultz is None:
        raise SystemExit("FAIL: Draft 1 anchor 2 (Schultz paragraph) not found")
    append_inline_ins(p_schultz, DRAFT1_INTRO_SCHULTZ_APPEND)
    print("OK Draft 1 (Intro applied-literature sentence appended)")

    # ---- Draft 2: Intro new paragraph before "We propose..." ----
    p_propose = find_paragraph_by_prefix(doc, ANCHORS["draft2_p17_intro_We_propose"])
    if p_propose is None:
        raise SystemExit("FAIL: Draft 2 anchor not found")
    new_p2 = new_paragraph_with_ins(DRAFT2_NEW_PARAGRAPH, pPr_source=p_propose._p.find(_q("pPr")))
    p_propose._p.addprevious(new_p2)
    print("OK Draft 2 (Intro implicit-injunctive-norm paragraph inserted)")

    # ---- Draft 3: GD foundational citation refresh (inline) ----
    p_gd_fit = find_paragraph_by_prefix(doc, ANCHORS["draft3_p170_GD_FIT"])
    if p_gd_fit is None:
        raise SystemExit("FAIL: Draft 3 anchor not found")
    if not replace_text_in_paragraph(p_gd_fit, DRAFT3_GD_CITATION_FIND, DRAFT3_GD_CITATION_REPLACE):
        raise SystemExit("FAIL: Draft 3 inline replace failed")
    print("OK Draft 3 (GD foundational citations refreshed)")

    # ---- Draft 4 (+ 5B as final sentence): GD two-pathway new paragraph ----
    # Insert after the GD-FIT paragraph (Draft 3 anchor).
    new_p4 = new_paragraph_with_ins(DRAFT4_NEW_PARAGRAPH, pPr_source=p_gd_fit._p.find(_q("pPr")))
    p_gd_fit._p.addnext(new_p4)
    print("OK Draft 4 (GD two-pathway paragraph inserted, with Draft 5B sentence)")

    # ---- Draft 5A: append political-ideology hedge to Study 4 Discussion ----
    p_pol = find_paragraph_by_prefix(doc, ANCHORS["draft5A_p132_study4_pol_ideology"])
    if p_pol is None:
        raise SystemExit("FAIL: Draft 5A anchor not found")
    append_inline_ins(p_pol, DRAFT5A_APPEND)
    print("OK Draft 5A (Study 4 Discussion political-ideology hedge appended)")

    # ---- Draft 6: Study 1 Discussion new paragraph ----
    p_s1 = find_paragraph_by_prefix(doc, ANCHORS["draft6_p50_study1_discussion"])
    if p_s1 is None:
        raise SystemExit("FAIL: Draft 6 anchor not found")
    new_p6 = new_paragraph_with_ins(DRAFT6_NEW_PARAGRAPH, pPr_source=p_s1._p.find(_q("pPr")))
    p_s1._p.addnext(new_p6)
    print("OK Draft 6 (Study 1 Discussion NPR-importance paragraph inserted)")

    # ---- Draft 7: References (alphabetical insertion) ----
    # Find every "References" heading paragraph, then walk forward to find the
    # right alphabetical slot for each new entry. Use the last-name first token
    # (e.g., "Andor") as the sort key.
    refs_heading = None
    for p in doc.paragraphs:
        if (p.text or "").strip() == "References":
            refs_heading = p
            break
    if refs_heading is None:
        raise SystemExit("FAIL: Draft 7: References heading not found")

    # Collect reference paragraphs (everything after the heading up to end of body
    # OR until a paragraph that looks like a non-reference body line). Many
    # APA-style ref lists are simply consecutive paragraphs to end of document.
    ref_paragraphs = []
    after_heading = False
    for p in doc.paragraphs:
        if p is refs_heading:
            after_heading = True
            continue
        if after_heading and (p.text or "").strip():
            ref_paragraphs.append(p)

    def sort_key(s: str) -> str:
        return s.split(",", 1)[0].strip().lower()

    for new_entry in DRAFT7_NEW_REFERENCES:
        new_key = sort_key(new_entry)
        # Insert before the first existing entry whose key sorts after new_key.
        slot = None
        for p in ref_paragraphs:
            existing_key = sort_key(p.text)
            if existing_key > new_key:
                slot = p
                break
        new_p = new_paragraph_with_ins(new_entry, pPr_source=(ref_paragraphs[0]._p.find(_q("pPr")) if ref_paragraphs else None))
        if slot is None:
            # Append at end (after last ref paragraph).
            if ref_paragraphs:
                ref_paragraphs[-1]._p.addnext(new_p)
            else:
                refs_heading._p.addnext(new_p)
        else:
            slot._p.addprevious(new_p)
        # Refresh ref_paragraphs list ordering (linear walk; cheap given ~50 refs).
        ref_paragraphs = []
        after_heading = False
        for p in doc.paragraphs:
            if p is refs_heading:
                after_heading = True
                continue
            if after_heading and (p.text or "").strip():
                ref_paragraphs.append(p)
    print(f"OK Draft 7 ({len(DRAFT7_NEW_REFERENCES)} References inserted alphabetically)")


# --- Post-save validation ---------------------------------------------------


def validate_output() -> None:
    if not OUT_DOCX.exists() or OUT_DOCX.stat().st_size == 0:
        raise SystemExit(f"FAIL: output missing/empty: {OUT_DOCX}")

    # Re-open and count w:ins elements with the expected author.
    doc2 = Document(str(OUT_DOCX))
    body = doc2.element.body
    inserts = body.findall(".//" + _q("ins"))
    deletes = body.findall(".//" + _q("del"))

    matching_author = [e for e in inserts if e.get(_q("author")) == AUTHOR]
    if len(matching_author) < 12:  # rough lower bound: 2 inline + 1 + 1 inline + 1 + 1 + 1 + 9 refs = 16+
        raise SystemExit(
            f"FAIL: only {len(matching_author)} w:ins elements with author={AUTHOR}; expected >= 12"
        )

    # Voice check: no en/em dashes in inserted text.
    bad_dashes = 0
    for ins in matching_author:
        for t in ins.findall(".//" + _q("t")):
            if t.text and ("–" in t.text or "—" in t.text):
                bad_dashes += 1
    if bad_dashes:
        raise SystemExit(f"FAIL: {bad_dashes} inserted text nodes contain em/en dashes")

    print(f"OK output: {OUT_DOCX} ({OUT_DOCX.stat().st_size} bytes)")
    print(f"OK w:ins count (author={AUTHOR}): {len(matching_author)}")
    print(f"OK w:del count total: {len(deletes)}")


# --- Main -------------------------------------------------------------------


def main() -> int:
    print("=== build_r2_tracked_drafts.py ===")
    validate_inputs()
    print(f"OK inputs present")
    validate_anchors()
    print("OK anchors present in manuscript_gdoc_text.txt")
    validate_no_dashes_in_drafts()
    print("OK draft prose passes voice constraints (no em/en dashes, no 'latent')")

    # Pre-delete output.
    if OUT_DOCX.exists():
        OUT_DOCX.unlink()
        print(f"Deleted prior output: {OUT_DOCX}")

    shutil.copyfile(SRC_DOCX, OUT_DOCX)
    print(f"Copied source: {SRC_DOCX} -> {OUT_DOCX}")

    doc = Document(str(OUT_DOCX))
    apply_drafts(doc)
    doc.save(str(OUT_DOCX))
    print(f"Saved: {OUT_DOCX}")

    validate_output()
    print("=== BUILD OK ===")
    return 0


if __name__ == "__main__":
    sys.exit(main())
